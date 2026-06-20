"""
Module 2 — Journal Entry Drafting

User describes a transaction in plain English. Optionally uploads a Chart of Accounts.
AI drafts a balanced double-entry journal entry using the company's actual account codes.

Reuses from Module 1:
  - get_chain()               : AI fallback chain (ai_client.py)
  - parse_coa()               : COA file parser (file_parser.py)
  - Fragment pattern          : @st.fragment at module level, st.rerun(scope="fragment")
  - Upfront spinner pattern   : check session_state before blocking work
  - Word doc builder pattern  : build_word_doc() adapted for journal entry output
"""

import re
import json
import io
from datetime import date
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.ai_client import get_chain
from utils.file_parser import parse_coa


# ── Transaction library ───────────────────────────────────────────────────

TRANSACTION_EXAMPLES = [
    ("Accrued salaries",        "Accrue unpaid salaries at month-end. Total amount is $85,000."),
    ("Prepaid insurance",       "Paid $24,000 for a 12-month insurance policy starting today."),
    ("Customer invoice",        "Issued an invoice to Acme Corp for $120,000 of consulting services rendered."),
    ("Cash receipt",            "Received $75,000 cash from Acme Corp against their outstanding invoice."),
    ("Monthly depreciation",    "Record monthly depreciation: equipment $8,500, vehicles $3,200."),
    ("Loan repayment",          "Made a loan repayment of $10,650 — $10,000 principal and $650 interest."),
    ("Deferred revenue",        "Recognised $30,000 of revenue that was previously deferred from a client deposit."),
    ("Inventory purchase",      "Purchased $45,000 of raw materials inventory on credit from a supplier."),
]

ACCOUNTING_STANDARDS = ["IFRS", "US GAAP", "Not sure"]


# ── AI prompt ─────────────────────────────────────────────────────────────

_JE_SYSTEM = """You are a senior accountant. Draft a balanced double-entry journal entry.

Rules:
- Total debits MUST equal total credits exactly.
- Use clear, professional account names.
- If a chart of accounts is provided, use ONLY account codes and names from that list.
- If no chart of accounts is provided, use standard account names and suggest typical codes.
- Apply the specified accounting standard where relevant (e.g. lease treatment, revenue recognition).
- The treatment_note must explain the accounting logic briefly — which standard or principle applies.

Return ONLY valid JSON in this exact format:
{
  "reference": "JE-001",
  "narration": "One-sentence description of the transaction",
  "lines": [
    {"account_code": "1001", "account_name": "Cash", "debit": 50000, "credit": 0},
    {"account_code": "2400", "account_name": "Deferred Revenue", "debit": 0, "credit": 50000}
  ],
  "treatment_note": "Brief explanation of accounting treatment and standard applied."
}

Use 0 (not null) for the side that has no amount. Do not include currency symbols in amounts."""


def _build_je_prompt(description: str, standard: str, company: str, coa_df) -> str:
    parts = []

    if company:
        parts.append(f"Company: {company}")

    parts.append(f"Accounting standard: {standard}")
    parts.append(f"Transaction: {description}")

    if coa_df is not None and not coa_df.empty:
        # Send up to 80 rows of the COA — enough context without wasting tokens
        sample = coa_df.head(80)
        rows = []
        for _, row in sample.iterrows():
            code = row.get("Account Code", "")
            name = row.get("Account Name", "")
            atype = row.get("Account Type", "")
            rows.append(f"{code} | {name} | {atype}" if atype else f"{code} | {name}")
        parts.append("\nChart of accounts (Code | Name | Type):\n" + "\n".join(rows))
        parts.append("\nIMPORTANT: Use ONLY account codes and names from the list above.")
    else:
        parts.append("\nNo chart of accounts provided. Use standard account names and suggest typical codes.")

    return "\n".join(parts)


# ── Word doc builder ──────────────────────────────────────────────────────

def _build_word_doc(je: dict, company: str, standard: str, model_name: str) -> bytes:
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Font defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Journal Entry")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x27, 0x44)

    # Metadata block
    doc.add_paragraph()
    meta_lines = [
        f"Reference:   {je.get('reference', 'JE-001')}",
        f"Date:        {date.today().strftime('%B %d, %Y')}",
        f"Standard:    {standard}",
    ]
    if company:
        meta_lines.insert(0, f"Company:     {company}")
    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # Narration
    nar = doc.add_paragraph()
    nar.add_run("Narration: ").bold = True
    nar.add_run(je.get("narration", ""))

    doc.add_paragraph()

    # Journal entry table
    lines = je.get("lines", [])
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, label in enumerate(["Account Code", "Account Name", "Debit ($)", "Credit ($)"]):
        hdr_cells[i].text = label
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)

    # Data rows
    total_debit = total_credit = 0.0
    for line in lines:
        row_cells = tbl.add_row().cells
        row_cells[0].text = str(line.get("account_code", ""))
        row_cells[1].text = str(line.get("account_name", ""))
        debit  = float(line.get("debit",  0) or 0)
        credit = float(line.get("credit", 0) or 0)
        row_cells[2].text = f"{debit:,.2f}"  if debit  else ""
        row_cells[3].text = f"{credit:,.2f}" if credit else ""
        total_debit  += debit
        total_credit += credit
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Totals row
    tot_cells = tbl.add_row().cells
    tot_cells[0].text = ""
    tot_cells[1].text = "TOTAL"
    tot_cells[2].text = f"{total_debit:,.2f}"
    tot_cells[3].text = f"{total_credit:,.2f}"
    for cell in tot_cells:
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(10)

    doc.add_paragraph()

    # Treatment note
    note_para = doc.add_paragraph()
    note_para.add_run("Accounting Treatment: ").bold = True
    note_para.add_run(je.get("treatment_note", ""))

    doc.add_paragraph()

    # Disclaimer footer
    disc = doc.add_paragraph(
        f"Generated by {model_name}. AI output — verify all account codes, amounts, "
        "and accounting treatments before posting to your accounting system."
    )
    disc.runs[0].font.size = Pt(9)
    disc.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Helpers ───────────────────────────────────────────────────────────────

def _fmt_amount(val) -> str:
    try:
        v = float(val or 0)
        return f"{v:,.2f}" if v else ""
    except (TypeError, ValueError):
        return ""


def _je_as_text(je: dict, company: str, standard: str, model_name: str) -> str:
    """Plain text version for copy-to-clipboard."""
    lines = [
        f"JOURNAL ENTRY — {je.get('reference', 'JE-001')}",
        f"Date: {date.today().strftime('%B %d, %Y')}",
    ]
    if company:
        lines.append(f"Company: {company}")
    lines += [
        f"Standard: {standard}",
        "",
        f"Narration: {je.get('narration', '')}",
        "",
        f"{'Account Code':<16} {'Account Name':<35} {'Debit':>14} {'Credit':>14}",
        "-" * 82,
    ]
    total_dr = total_cr = 0.0
    for line in je.get("lines", []):
        dr = float(line.get("debit",  0) or 0)
        cr = float(line.get("credit", 0) or 0)
        total_dr += dr
        total_cr += cr
        lines.append(
            f"{str(line.get('account_code','')):<16} "
            f"{str(line.get('account_name','')):<35} "
            f"{(f'{dr:,.2f}' if dr else '')!s:>14} "
            f"{(f'{cr:,.2f}' if cr else '')!s:>14}"
        )
    lines += [
        "-" * 82,
        f"{'':16} {'TOTAL':<35} {f'{total_dr:,.2f}':>14} {f'{total_cr:,.2f}':>14}",
        "",
        f"Treatment: {je.get('treatment_note', '')}",
        "",
        f"Generated by {model_name}. AI output — verify before posting.",
    ]
    return "\n".join(lines)


# ── Fragments ─────────────────────────────────────────────────────────────
# Defined at module level — never inside render() — per Streamlit Build Rule 7.

@st.fragment
def _coa_section():
    """COA file upload. Parses and caches the chart of accounts."""
    st.markdown("#### Chart of Accounts (optional)")
    st.caption(
        "Upload your chart of accounts so the AI uses your actual account codes and names. "
        "Accepted formats: .xlsx, .xls, .csv. "
        "File must have at least an account code column and an account name column."
    )

    uploaded = st.file_uploader(
        "Upload COA file",
        type=["xlsx", "xls", "csv"],
        key="je_coa_upload",
        label_visibility="collapsed",
    )

    if uploaded:
        cache_key = f"je_coa_{uploaded.name}"
        if cache_key not in st.session_state:
            with st.spinner("Reading chart of accounts..."):
                try:
                    chain = get_chain(st.session_state)
                    coa_df = parse_coa(uploaded.read(), uploaded.name, chain)
                    st.session_state[cache_key]        = coa_df
                    st.session_state["je_coa_filename"] = uploaded.name
                except Exception as e:
                    st.error(f"Could not read COA file: {e}")
                    return

        coa_df = st.session_state.get(cache_key)
        if coa_df is not None:
            st.success(f"{len(coa_df):,} accounts loaded from {uploaded.name}")
            with st.expander("Preview accounts", expanded=False):
                st.dataframe(coa_df.head(20), use_container_width=True, hide_index=True)
    else:
        # Clear cached COA if user removed the file
        if "je_coa_filename" in st.session_state:
            old_key = f"je_coa_{st.session_state['je_coa_filename']}"
            st.session_state.pop(old_key, None)
            st.session_state.pop("je_coa_filename", None)
        st.caption("No COA loaded — AI will suggest standard account names.")


@st.fragment
def _je_section(standard: str, company: str):
    """Transaction input, Generate button, output display, and downloads."""

    st.markdown("#### Transaction")

    # Transaction library — click to pre-fill
    st.caption("No transaction in mind? Click an example to load it.")
    cols = st.columns(4)
    for i, (label, text) in enumerate(TRANSACTION_EXAMPLES):
        if cols[i % 4].button(label, key=f"ex_{i}", use_container_width=True):
            st.session_state["je_input"] = text
            st.rerun(scope="fragment")

    st.markdown("")

    description = st.text_area(
        "Describe the transaction",
        key="je_input",
        height=100,
        placeholder="e.g. Paid $24,000 for a 12-month insurance policy starting today.",
        help=(
            "Describe the transaction in plain English. Include the amount, the parties involved, "
            "and the purpose. The more detail you provide, the more accurate the entry."
        ),
    )

    if st.button("Generate Journal Entry", type="primary", key="je_generate"):
        if not description or not description.strip():
            st.warning("Please describe a transaction before generating.")
            st.stop()

        # Get COA if loaded
        coa_filename = st.session_state.get("je_coa_filename")
        coa_df = st.session_state.get(f"je_coa_{coa_filename}") if coa_filename else None

        with st.spinner("Drafting journal entry..."):
            try:
                chain = get_chain(st.session_state)
                prompt = _build_je_prompt(description, standard, company, coa_df)
                messages = [
                    {"role": "system", "content": _JE_SYSTEM},
                    {"role": "user",   "content": prompt},
                ]
                response, model_name = chain.complete(messages, timeout=30)
                clean = re.sub(r"```(?:json)?|```", "", response).strip()
                match = re.search(r"\{.*\}", clean, re.DOTALL)
                if not match:
                    raise ValueError("AI did not return a valid journal entry. Please try again.")
                je = json.loads(match.group(0))

                # Validate balance
                total_dr = sum(float(l.get("debit",  0) or 0) for l in je.get("lines", []))
                total_cr = sum(float(l.get("credit", 0) or 0) for l in je.get("lines", []))
                if abs(total_dr - total_cr) > 0.01:
                    je["treatment_note"] = (
                        je.get("treatment_note", "") +
                        f" [Note: entry may not balance — debit {total_dr:,.2f} vs credit {total_cr:,.2f}. Review before posting.]"
                    )

                st.session_state["je_result"]     = je
                st.session_state["je_model"]      = model_name
                st.session_state["je_word_bytes"] = _build_word_doc(je, company, standard, model_name)
                st.session_state["je_text"]       = _je_as_text(je, company, standard, model_name)

            except Exception as e:
                st.error(f"Could not generate journal entry: {e}")
                return

    # ── Output display ─────────────────────────────────────────────────────
    je         = st.session_state.get("je_result")
    model_name = st.session_state.get("je_model", "")
    word_bytes = st.session_state.get("je_word_bytes")
    je_text    = st.session_state.get("je_text", "")

    if je:
        st.markdown("---")

        # Header
        ref  = je.get("reference", "JE-001")
        nar  = je.get("narration", "")
        today = date.today().strftime("%B %d, %Y")

        col_ref, col_date = st.columns([1, 1])
        col_ref.markdown(f"**{ref}**")
        col_date.markdown(f"*{today}*")
        st.markdown(f"**Narration:** {nar}")
        st.markdown("")

        # Entry table
        lines = je.get("lines", [])
        total_dr = total_cr = 0.0

        header_cols = st.columns([2, 4, 2, 2])
        header_cols[0].markdown("**Account Code**")
        header_cols[1].markdown("**Account Name**")
        header_cols[2].markdown("**Debit ($)**")
        header_cols[3].markdown("**Credit ($)**")
        st.divider()

        for line in lines:
            dr = float(line.get("debit",  0) or 0)
            cr = float(line.get("credit", 0) or 0)
            total_dr += dr
            total_cr += cr
            row_cols = st.columns([2, 4, 2, 2])
            row_cols[0].write(str(line.get("account_code", "")))
            row_cols[1].write(str(line.get("account_name", "")))
            row_cols[2].write(f"{dr:,.2f}" if dr else "")
            row_cols[3].write(f"{cr:,.2f}" if cr else "")

        st.divider()
        tot_cols = st.columns([2, 4, 2, 2])
        tot_cols[1].markdown("**TOTAL**")
        tot_cols[2].markdown(f"**{total_dr:,.2f}**")
        tot_cols[3].markdown(f"**{total_cr:,.2f}**")

        # Balance check
        if abs(total_dr - total_cr) <= 0.01:
            st.success("Entry balances. Debits = Credits.")
        else:
            st.error(f"Entry does not balance. Debit {total_dr:,.2f} ≠ Credit {total_cr:,.2f}. Review before posting.")

        # Treatment note
        st.markdown("")
        st.markdown(f"**Accounting Treatment:** {je.get('treatment_note', '')}")

        # Model name + disclaimer
        short_model = model_name.split("/")[-1] if "/" in model_name else model_name
        st.caption(
            f"Generated by {short_model} · "
            "AI output — review all account codes, amounts, and treatments before posting."
        )

        # Downloads
        st.markdown("")
        dl_col, copy_col, _ = st.columns([2, 2, 4])

        filename = f"JE_{date.today().strftime('%Y-%m-%d')}.docx"
        dl_col.download_button(
            label="Download Word doc",
            data=word_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="je_download",
        )

        copy_col.text_area(
            "Copy to clipboard",
            value=je_text,
            height=68,
            key="je_copy_area",
            help="Select all text here (Ctrl+A) and copy to paste into your accounting system.",
        )


# ── Learning Guide ────────────────────────────────────────────────────────

def _learning_guide():
    with st.expander("Learning Guide — Journal Entries", expanded=False):
        st.markdown("""
**What is a journal entry?**

A journal entry is the fundamental record of a financial transaction.
Every event that affects a company's finances — a payment, an invoice, a depreciation charge — is first recorded as a journal entry before it flows into the financial statements.

**Double-entry bookkeeping**

Every journal entry has two sides: debits and credits. They must always balance.
- **Debit** increases assets and expenses. It decreases liabilities, equity, and revenue.
- **Credit** increases liabilities, equity, and revenue. It decreases assets and expenses.

A simple rule: for every debit, there must be an equal credit somewhere.

**IFRS vs US GAAP**

Both frameworks follow the same double-entry principles. The differences appear in specific areas:
- **Leases (IFRS 16 vs ASC 842):** Both now require most leases on the balance sheet, but the income statement treatment differs.
- **Revenue recognition (IFRS 15 vs ASC 606):** Both use a five-step model. Minor differences in application.
- **Inventory (IFRS prohibits LIFO. US GAAP allows it).**

If you select "Not sure," the AI notes where the treatment might differ.

**Chart of accounts**

A chart of accounts is the numbered list of all accounts a company uses — Cash, Accounts Receivable, Revenue, Salaries Expense, etc. Each account has a unique code. Uploading yours means the AI uses your exact codes instead of generic ones.

**Before posting**

AI-generated entries are a starting point. Always verify:
- Account codes match your accounting system
- Amounts are correct
- The accounting treatment is appropriate for your jurisdiction and policy
""")


# ── Render ────────────────────────────────────────────────────────────────

def render():
    st.markdown("## Journal Entry Drafting")
    st.caption(
        "Describe a transaction in plain English. "
        "Upload your chart of accounts for entries that match your account codes. "
        "AI drafts a balanced double-entry journal entry ready for review."
    )
    st.markdown("---")

    # ── Global settings ────────────────────────────────────────────────────
    col_std, col_co, _ = st.columns([2, 2, 4])

    standard = col_std.selectbox(
        "Accounting standard",
        options=ACCOUNTING_STANDARDS,
        key="je_standard",
        help=(
            "IFRS is used in most countries outside the US. "
            "US GAAP applies to US-based companies. "
            "If unsure, select 'Not sure' and the AI will note where treatments differ."
        ),
    )

    company = col_co.text_input(
        "Company name (optional)",
        key="je_company",
        placeholder="e.g. Acme Corp",
        help="Used to personalise the journal entry header and narration.",
    )

    st.markdown("---")

    # ── COA section (fragment) ─────────────────────────────────────────────
    _coa_section()

    st.markdown("---")

    # ── Journal entry section (fragment) ──────────────────────────────────
    _je_section(standard, company)

    st.markdown("---")

    # ── Learning Guide ─────────────────────────────────────────────────────
    _learning_guide()
