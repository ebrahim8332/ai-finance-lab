"""
Module 1: Variance Commentary

User uploads a budget vs. actual Excel file. App auto-detects company name and
period from file content. Shows data table and five finance charts automatically.
AI generates structured commentary tailored to the chosen audience.
Output is collapsible. Download as a Word document with charts embedded.
"""

import io
import re
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import pandas as pd

from utils.ai_client import get_chain
from utils.excel_parser import dataframe_to_text, extract_file_metadata
from utils.file_parser import detect_structure, extract_data
from utils.chart_builder import build_all_charts, fig_to_png_bytes, prepare_data
from utils.base import AllProvidersExhausted


def _safe_md(text: str) -> str:
    """Escapes dollar signs so Streamlit doesn't render them as LaTeX math."""
    return text.replace("$", r"\$")


def _build_display_df(df: pd.DataFrame) -> pd.DataFrame:
    """
    Returns a formatted copy of the dataframe for screen display only.
    The raw df (unmodified) goes to the AI and the chart builder.

    Rules applied:
    - Dollar / number columns (budget, actual, variance): comma-separated integers  e.g. 1,234
    - Variance % columns: multiplied by 100 and shown as  e.g. 5.7%
    - All other columns: left as-is
    """
    disp = df.copy()
    for col in disp.columns:
        col_lower = col.lower()
        if '%' in col_lower or 'pct' in col_lower or 'percent' in col_lower:
            disp[col] = disp[col].apply(
                lambda v: f"{v * 100:.1f}%"
                if pd.notna(v) and isinstance(v, (int, float))
                else (str(v) if pd.notna(v) else "")
            )
        elif any(kw in col_lower for kw in ['budget', 'actual', 'variance', '$', '($k)', 'amount']):
            disp[col] = disp[col].apply(
                lambda v: f"{v:,.0f}"
                if pd.notna(v) and isinstance(v, (int, float))
                else (str(v) if pd.notna(v) else "")
            )
    return disp


def _build_materiality_df(df: pd.DataFrame, pct_threshold: float, amt_threshold: float) -> pd.DataFrame:
    """
    Adds a Status column to the display dataframe.

    An item is material when BOTH conditions are true:
      - absolute variance >= amt_threshold
      - absolute variance % >= pct_threshold

    Status values:
      🔴 Unfavorable  — material and over budget (bad)
      🟢 Favorable    — material and under budget (good)
      ⚪ Immaterial   — below threshold, not worth investigating
    """
    _, cols = prepare_data(df)
    budget_col = cols.get("budget")
    actual_col = cols.get("actual")

    disp = _build_display_df(df)

    if not budget_col or not actual_col:
        return disp

    statuses = []
    for _, row in df.iterrows():
        try:
            budget = float(row[budget_col])
            actual = float(row[actual_col])
            variance = actual - budget
            pct = abs(variance / budget) * 100 if budget != 0 else 0
            material = abs(variance) >= amt_threshold and pct >= pct_threshold
            if not material:
                statuses.append("⚪ Immaterial")
            elif variance < 0:
                statuses.append("🔴 Unfavorable")
            else:
                statuses.append("🟢 Favorable")
        except (TypeError, ValueError):
            statuses.append("")

    disp.insert(len(disp.columns), "Status", statuses)
    return disp


AUDIENCE_DESCRIPTIONS = {
    "CFO":        "senior finance executive who wants concise, technically precise commentary "
                  "with focus on margin drivers, working capital, and corrective actions",
    "Board":      "board of directors who want high-level narrative, key risks, and strategic "
                  "implications — minimal technical detail, maximum clarity",
    "Operations": "operational managers who need to understand what drove variances in their "
                  "areas and what actions they should take this quarter",
}

AUDIENCE_TONE = {
    "CFO":        "concise, technically precise, focused on root cause and action",
    "Board":      "strategic, narrative-driven, focused on implications and risk",
    "Operations": "practical, direct, focused on what happened and what to do next",
}


# ── Word document builder ─────────────────────────────────────────────────

def _add_bold_runs(paragraph, text: str):
    """Converts **bold** markdown markers to actual bold formatting in Word."""
    for part in re.split(r'(\*\*.*?\*\*)', text):
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def build_word_doc(
    commentary_text: str,
    company: str,
    period: str,
    audience: str,
    chart_pngs: list[tuple[str, bytes]],  # list of (chart_title, png_bytes)
) -> bytes:
    """
    Builds a Word document with charts embedded above the AI commentary.
    Returns bytes for the Streamlit download button.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    run = title.add_run(f'{company} — {period} Management Commentary')
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f'Audience: {audience}')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    sub.runs[0].font.size = Pt(10)

    # ── Charts section ────────────────────────────────────────────────────
    if chart_pngs:
        doc.add_paragraph()
        charts_heading = doc.add_paragraph()
        charts_heading.add_run('Charts').bold = True
        charts_heading.runs[0].font.size = Pt(14)

        for chart_title, png_bytes in chart_pngs:
            if png_bytes:
                doc.add_paragraph(chart_title).runs[0].italic = True
                doc.add_picture(io.BytesIO(png_bytes), width=Inches(6.0))
                doc.add_paragraph()  # spacing between charts

    # ── Commentary section ────────────────────────────────────────────────
    doc.add_paragraph()
    comm_heading = doc.add_paragraph()
    comm_heading.add_run('Management Commentary').bold = True
    comm_heading.runs[0].font.size = Pt(14)
    doc.add_paragraph()

    for line in commentary_text.splitlines():
        s = line.strip()
        if not s:
            doc.add_paragraph()
        elif s.startswith('### '):
            doc.add_heading(s[4:], level=3)
        elif s.startswith('## '):
            doc.add_heading(s[3:], level=2)
        elif s.startswith('# '):
            doc.add_heading(s[2:], level=1)
        elif s.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            _add_bold_runs(p, s[2:])
        else:
            p = doc.add_paragraph()
            _add_bold_runs(p, s)

    doc.add_paragraph()
    footer = doc.add_paragraph('AI-generated content. Review all figures before distributing.')
    footer.runs[0].italic = True
    footer.runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Prompt builder ────────────────────────────────────────────────────────

def build_prompt(data_text: str, audience: str, period: str, company: str) -> list[dict]:
    system = f"""You are a senior FP&A analyst writing management commentary for {company}.
Your audience is a {AUDIENCE_DESCRIPTIONS[audience]}.
Tone: {AUDIENCE_TONE[audience]}.

Rules:
- Write in clear, professional English. No jargon for the sake of it.
- Do not use phrases like "leverage", "ecosystem", "transformative", or "actionable insights".
- Label every section clearly with a markdown heading (## Section Name).
- Be specific: reference actual numbers from the data.
- All figures are in USD thousands unless stated otherwise.
- Do not hallucinate. Only reference line items present in the data.
- End with a Conclusions section that adds new synthesis, not a restatement.
"""
    user = f"""Write management commentary for {company} for the period: {period}.

FINANCIAL DATA:
{data_text}

Structure the commentary as follows:
## Executive Summary
## Revenue Analysis
## Cost and Margin Analysis
## Operating Performance
## Key Risks and Watch Items
## Conclusions

Audience: {audience}. Tailor depth and language accordingly.
Write in paragraphs, not bullet points.
"""
    return [
        {"role": "system", "content": system},
        {"role": "user",   "content": user},
    ]


# ── Fragment functions (module level — defined once, not on every render) ─
# Defining @st.fragment inside render() causes Streamlit to re-register the
# fragment on every rerun, which triggers additional reruns and creates a
# gray-screen cycle. Defining them here at module level fixes that.

@st.fragment
def _data_table_section(df):
    st.markdown("**Data loaded from file:**")

    mat_col1, mat_col2 = st.columns([1, 1])
    with mat_col1:
        pct_threshold = st.slider(
            "Materiality threshold — variance %",
            min_value=1, max_value=20, value=5, step=1,
            help="An item must exceed both thresholds to be flagged as material.",
        )
    with mat_col2:
        amt_threshold = st.slider(
            "Materiality threshold — variance amount ($)",
            min_value=1000, max_value=100000, value=10000, step=1000,
            format="$%d",
            help="An item must exceed both thresholds to be flagged as material.",
        )

    st.dataframe(
        _build_materiality_df(df, pct_threshold, amt_threshold),
        use_container_width=True,
        hide_index=True,
    )
    st.caption(
        "🔴 Unfavorable = material and over budget.  "
        "🟢 Favorable = material and under budget.  "
        "⚪ Immaterial = below threshold, no action needed.  "
        "Both thresholds must be exceeded for an item to be flagged."
    )

    _, detected_cols = prepare_data(df)
    col_parts = []
    for role, label in [("Label", "label"), ("Budget", "budget"), ("Actual", "actual")]:
        val = detected_cols.get(label)
        if val and not val.startswith("_"):
            col_parts.append(f"**{role}:** {val}")
    if col_parts:
        st.caption("Detected columns — " + "  ·  ".join(col_parts) +
                   "  ·  *If any column is wrong, rename it in your file and re-upload.*")


@st.fragment
def _qa_section(df):
    if st.session_state.pop("_qa_clear_input", False):
        st.session_state["qa_input"] = ""

    st.markdown("#### Ask a question about your data")
    st.caption(
        "Ask anything: 'Which line item had the biggest variance?' · "
        "'What drove the drop in net income?' · "
        "'Which three items should management focus on?'"
    )

    qa_col1, qa_col2 = st.columns([4, 1])
    with qa_col1:
        question = st.text_input(
            "Your question",
            placeholder="e.g. Which expenses were most over budget?",
            label_visibility="collapsed",
            key="qa_input",
        )
    with qa_col2:
        ask_clicked = st.button("Ask AI", type="primary", use_container_width=True)

    if ask_clicked and question.strip():
        data_text = dataframe_to_text(df)
        qa_messages = [
            {
                "role": "system",
                "content": (
                    "You are a senior FP&A analyst. Answer the user's question using only "
                    "the financial data provided. Be specific — reference actual numbers. "
                    "Keep the answer concise: 3-5 sentences maximum. "
                    "Do not reference any figures not present in the data."
                ),
            },
            {
                "role": "user",
                "content": f"FINANCIAL DATA:\n{data_text}\n\nQUESTION: {question.strip()}",
            },
        ]
        with st.spinner("Thinking..."):
            try:
                st.session_state.pop("locked_provider_index", None)
                st.session_state.pop("_fallback_errors", None)
                chain = get_chain(st.session_state)
                answer, model_used_qa = chain.complete(qa_messages, timeout=60)
                if "qa_history" not in st.session_state:
                    st.session_state["qa_history"] = []
                st.session_state["qa_history"].append({
                    "question": question.strip(),
                    "answer":   answer,
                    "model":    model_used_qa,
                })
            except AllProvidersExhausted as e:
                st.error(str(e))
            except Exception as e:
                st.error(f"Unexpected error: {e}")

    if st.session_state.get("qa_history"):
        for item in reversed(st.session_state["qa_history"]):
            with st.expander(f"💬 {item['question']}", expanded=True):
                st.markdown(_safe_md(item["answer"]))
                st.caption(f"AI response from {item['model']}")
        if st.button("Clear questions", key="clear_qa"):
            st.session_state["qa_history"] = []
            st.session_state["_qa_clear_input"] = True
            st.rerun(scope="fragment")


@st.fragment
def _commentary_section(df, audience, period, company, charts):
    if st.button("Generate AI Commentary", type="primary", use_container_width=True):
        data_text = dataframe_to_text(df)
        messages  = build_prompt(data_text, audience, period, company)
        with st.spinner("Analysing variances and writing commentary..."):
            try:
                st.session_state.pop("locked_provider_index", None)
                st.session_state.pop("_fallback_errors", None)
                chain = get_chain(st.session_state)
                response, model_used = chain.complete(messages, timeout=120)
            except AllProvidersExhausted as e:
                st.error(str(e))
                return
            except Exception as e:
                st.error(f"Unexpected error: {e}")
                return

        st.session_state["var_commentary"] = response
        st.session_state["var_model_used"] = model_used
        st.session_state["var_company"]    = company
        st.session_state["var_period"]     = period
        st.session_state["var_audience"]   = audience
        st.session_state["var_charts"]     = charts

    if "var_commentary" in st.session_state:
        response     = st.session_state["var_commentary"]
        model_used   = st.session_state["var_model_used"]
        company_out  = st.session_state["var_company"]
        period_out   = st.session_state["var_period"]
        audience_out = st.session_state["var_audience"]
        saved_charts = st.session_state.get("var_charts", [])

        st.markdown(
            '<hr style="border: none; border-top: 3px solid #d0d0d0; margin: 18px 0 20px 0;">',
            unsafe_allow_html=True,
        )
        with st.expander(
            f"📋 AI Commentary — {company_out}, {period_out} ({audience_out})",
            expanded=True,
        ):
            st.markdown(_safe_md(response))
            st.caption(f"AI response from {model_used}")

        chart_pngs = [(title, fig_to_png_bytes(fig)) for title, fig in saved_charts]
        word_bytes = build_word_doc(response, company_out, period_out, audience_out, chart_pngs)
        st.download_button(
            label="⬇ Download as Word document (.docx)",
            data=word_bytes,
            file_name=f"{company_out.lower().replace(' ', '_')}_{period_out.lower().replace(' ', '_')}_commentary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        st.caption("Review all AI-generated content before distributing. Verify figures against source data.")


# ── Module renderer ───────────────────────────────────────────────────────

_FINANCE_NOISE = [
    "variance", "budget", "actual", "template", "analysis",
    "report", "p&l", "income", "expense", "fiscal", "period",
    "management", "accounts", "statement",
]


def _process_file(file_bytes: bytes, filename: str) -> dict | None:
    """
    Runs all expensive AI work for a newly uploaded file and stores results
    in session_state under a key derived from the filename.

    Returns the cached result dict, or None if processing failed.
    Called only when a new file is uploaded — all reruns read from cache.
    """
    cache_key = f"file_data_{filename}"
    if cache_key in st.session_state:
        return st.session_state[cache_key]

    st.session_state.pop("locked_provider_index", None)
    st.session_state.pop("_fallback_errors", None)
    chain = get_chain(st.session_state)

    # Step 1: detect structure (AI call — ~50 tokens)
    try:
        structure = detect_structure(file_bytes, chain)
    except Exception as e:
        st.error(f"Could not detect file structure: {e}")
        return None

    # Step 2: extract dataframe (no AI — uses column indices from step 1)
    # For horizontal files, default to the first period; user can change it later.
    selected_period = None
    if structure["format"] == "horizontal" and structure.get("period_names"):
        selected_period = structure["period_names"][0]

    try:
        df = extract_data(file_bytes, structure, period=selected_period)
    except Exception as e:
        st.error(
            f"Could not extract data: {e} "
            "Try simplifying the file to four columns: Line Item, Budget, Actual, Variance %."
        )
        return None

    # Step 3: build charts (AI call for line item grouping — ~50-80 tokens)
    chart_result = build_all_charts(df, chain)

    result = {
        "structure":       structure,
        "df":              df,
        "charts":          chart_result["charts"],
        "file_type":       chart_result["file_type"],
        "selected_period": selected_period,
    }
    st.session_state[cache_key] = result
    return result


def render():
    st.markdown("## Variance Commentary")
    st.markdown(
        "Upload a budget vs. actual Excel file. "
        "The app reads the company name and reporting period directly from the file, "
        "generates finance charts automatically, and writes AI commentary "
        "tailored to your chosen audience. Download everything as a Word document."
    )

    st.markdown(
        '<hr style="border: none; border-top: 3px solid #d0d0d0; margin: 18px 0 20px 0;">',
        unsafe_allow_html=True,
    )

    # ── Inputs ────────────────────────────────────────────────────────────
    col1, col2 = st.columns([2, 1])

    with col1:
        uploaded_file = st.file_uploader(
            "Upload your budget vs. actual file",
            type=["xlsx", "xls"],
            help="Works with most Excel variance reports. Data should start in row 1 "
                 "with column headers. If the preview looks wrong, move your data to "
                 "the first sheet and remove any title rows above the headers.",
        )

    if uploaded_file is None:
        with col2:
            st.selectbox("Commentary audience", options=["CFO", "Board", "Operations"],
                         help="Changes the tone, depth, and focus of the commentary.")
            st.text_input("Company name", placeholder="e.g. Acme Corp",
                          help="Auto-detected from the file. Edit if incorrect.")
            st.text_input("Reporting period", placeholder="e.g. Q2 2026",
                          help="Auto-detected from the file. Edit if incorrect.")
        return

    # ── File processing (runs once per file, all results cached) ─────────
    # Check if this is a new file using only session_state — no blocking work yet.
    # If it is new, show the spinner BEFORE reading bytes or calling AI so the
    # user sees feedback immediately instead of a blank screen.
    cache_key = f"file_data_{uploaded_file.name}"
    is_new_file = cache_key not in st.session_state

    if is_new_file:
        with st.spinner("Analysing file — detecting structure, extracting data, building charts..."):
            try:
                file_bytes = uploaded_file.read()
                uploaded_file.seek(0)
            except Exception as e:
                st.error(f"Could not read the file: {e}")
                return

            meta = extract_file_metadata(uploaded_file)
            raw_company = meta["company"] or ""
            _is_noise = any(kw in raw_company.lower() for kw in _FINANCE_NOISE)
            st.session_state["detected_company"] = (
                raw_company if raw_company and not _is_noise and len(raw_company) <= 60 else ""
            )
            st.session_state["detected_period"] = meta["period"] or ""
            st.session_state["_last_uploaded"]  = uploaded_file.name

            cached = _process_file(file_bytes, uploaded_file.name)
            if cached is None:
                return
    else:
        try:
            file_bytes = uploaded_file.read()
            uploaded_file.seek(0)
        except Exception as e:
            st.error(f"Could not read the file: {e}")
            return
        cached = st.session_state[cache_key]

    structure       = cached["structure"]
    df              = cached["df"]
    charts          = cached["charts"]
    file_type       = cached.get("file_type", "pl")
    selected_period = cached["selected_period"]

    # ── Audience / company / period inputs ────────────────────────────────
    with col2:
        audience = st.selectbox(
            "Commentary audience",
            options=["CFO", "Board", "Operations"],
            help="Changes the tone, depth, and focus of the commentary.",
        )
        company = st.text_input(
            "Company name",
            value=st.session_state.get("detected_company", ""),
            placeholder="e.g. Acme Corp",
            help="Auto-detected from the file. Edit if incorrect.",
        )
        period = st.text_input(
            "Reporting period",
            value=st.session_state.get("detected_period", "Current Period"),
            placeholder="e.g. Q2 2026",
            help="Auto-detected from the file. Edit if incorrect.",
        )

    st.markdown(
        '<hr style="border: none; border-top: 3px solid #d0d0d0; margin: 18px 0 20px 0;">',
        unsafe_allow_html=True,
    )

    # ── Period selector for horizontal files ─────────────────────────────
    if structure["format"] == "horizontal" and structure.get("period_names"):
        st.info("This file has a monthly layout. Select the period you want to analyse.")
        chosen_period = st.selectbox(
            "Select period",
            options=structure["period_names"],
            index=structure["period_names"].index(selected_period) if selected_period else 0,
            key="selected_period",
        )
        # If the user picks a different period, clear the cache so data re-extracts
        if chosen_period != selected_period:
            del st.session_state[f"file_data_{uploaded_file.name}"]
            st.rerun()

    # ── Structure caption ─────────────────────────────────────────────────
    if structure["_source"] == "ai":
        st.caption(
            "Structure detected automatically "
            f"({'horizontal — ' + selected_period if selected_period else 'vertical'})."
        )
    else:
        st.caption("Structure detected using keyword matching (AI unavailable).")

    # ── Data table with materiality indicators ───────────────────────────
    _data_table_section(df)

    st.markdown(
        '<hr style="border: none; border-top: 3px solid #d0d0d0; margin: 18px 0 20px 0;">',
        unsafe_allow_html=True,
    )

    # ── Charts ────────────────────────────────────────────────────────────
    if charts:
        with st.expander("📊 Charts", expanded=True):
            for i in range(0, len(charts), 2):
                cols_chart = st.columns(2)
                for j, (title, fig) in enumerate(charts[i:i+2]):
                    with cols_chart[j]:
                        with st.container(border=True):
                            st.plotly_chart(fig, use_container_width=True)

    st.markdown(
        '<hr style="border: none; border-top: 3px solid #d0d0d0; margin: 18px 0 20px 0;">',
        unsafe_allow_html=True,
    )

    # ── Learning Guide ────────────────────────────────────────────────────
    with st.expander("📖 Learning Guide — How to Read This Analysis", expanded=False):
        st.markdown("""
### What is Variance Analysis?

Variance analysis compares what was planned (Budget) against what actually happened (Actual).
The difference is the variance. Finance teams use it to identify where performance deviated
from plan and why — so management can take corrective action.

A variance is not automatically good or bad. Context matters:
- Revenue over budget is favorable. Revenue under budget is unfavorable.
- Costs under budget is favorable. Costs over budget is unfavorable.

---

### What is Materiality?

Not every variance deserves management attention. A $500 overrun on a $10M cost base is noise.
Materiality is a threshold that filters out the noise.

Finance teams typically set two thresholds: a dollar amount and a percentage.
An item is only flagged when it exceeds **both**. This prevents small items with large percentages
(e.g. a $100 item that is 50% over budget) from cluttering the analysis.

The sliders above let you set your own thresholds. The default — 5% and $10,000 — is a common
starting point for mid-sized companies.

---

### How to Read the Waterfall Bridge Chart

The bridge chart shows how you got from Revenue Budget to Revenue Actual.
Each bar is a category (Revenue, COGS, Opex). The direction tells you the profit impact:
- A bar going **up** means that category improved profit (favorable)
- A bar going **down** means that category hurt profit (unfavorable)

Revenue bars go up when actual revenue exceeded budget. Cost bars go up when actual costs were
**below** budget (spending less than planned is good for profit).

This chart answers: **Where did the profit gap come from?**

---

### How to Read the Budget vs Actual Chart

This chart shows Budget and Actual side by side for each category.
It is the simplest comparison — no sign adjustments, just raw numbers.

Use it to quickly see which categories came in above or below plan across the board.

---

### How to Read the Top 5 Movers Chart

This chart shows the five line items with the biggest profit impact.
The x-axis shows profit impact, not raw variance:
- Revenue items: a negative actual-vs-budget variance moves profit left (unfavorable)
- Cost items: a positive actual-vs-budget variance moves profit left (unfavorable) because spending more hurts profit

The bar **label** shows the raw variance (Actual minus Budget) so you can see the actual numbers.
The bar **direction** shows whether that variance was good or bad for profit.

Green = favorable impact on profit. Red = unfavorable impact on profit.

This chart answers: **Which specific line items drove the result?**

---

### Why Does Audience Matter?

The same financial data tells a different story depending on who is reading it.

- **CFO:** Wants root cause, margin drivers, and specific corrective actions. Technical and precise.
- **Board:** Wants the strategic narrative, key risks, and implications. Minimal numbers, maximum clarity.
- **Operations:** Wants to know what happened in their area and what to do next. Practical and direct.

The AI commentary changes its structure, language, and level of detail based on the audience you select.
""")

    st.markdown(
        '<hr style="border: none; border-top: 3px solid #d0d0d0; margin: 18px 0 20px 0;">',
        unsafe_allow_html=True,
    )

    # ── AI Q&A ────────────────────────────────────────────────────────────
    _qa_section(df)

    st.markdown(
        '<hr style="border: none; border-top: 3px solid #d0d0d0; margin: 18px 0 20px 0;">',
        unsafe_allow_html=True,
    )

    # ── Generate Commentary ───────────────────────────────────────────────
    _commentary_section(df, audience, period, company, charts)
